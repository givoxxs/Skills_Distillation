from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, Mm
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DOCX_PATH = Path(__file__).parent / "paper_draft.docx"


def set_run_font(run, name="Times New Roman", size=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def paragraph_has_drawing(paragraph):
    return bool(paragraph._p.xpath(".//w:drawing"))


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def format_docx(path=DOCX_PATH):
    doc = Document(path)

    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for style_name in ["Normal", "Body Text", "First Paragraph", "Compact"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(10)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            style.paragraph_format.space_after = Pt(4)
            style.paragraph_format.line_spacing = 1.0

    for style_name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = None
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(4)

    if "Caption" in doc.styles:
        cap = doc.styles["Caption"]
        cap.font.name = "Times New Roman"
        cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        cap.font.size = Pt(9)
        cap.font.italic = False
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        cap.paragraph_format.space_before = Pt(2)

    in_abstract = False
    before_abstract = True
    for para in doc.paragraphs:
        text = para.text.strip()

        if paragraph_has_drawing(para):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if not text:
            continue

        for run in para.runs:
            if para.style and para.style.name.startswith("Heading"):
                set_run_font(run, size=None, bold=True)
            elif para.style and para.style.name == "Caption":
                set_run_font(run, size=9)
            else:
                set_run_font(run, size=10)

        if text == "Abstract":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            in_abstract = True
            before_abstract = False
            continue

        if text.startswith("Keywords:"):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            in_abstract = False
            before_abstract = False
            if para.runs:
                para.runs[0].font.bold = True
            continue

        if text.startswith("Figure "):
            para.style = (
                doc.styles["Caption"] if "Caption" in doc.styles else para.style
            )
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_font(run, size=9)
            continue

        if before_abstract:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if in_abstract:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            continue

        if para.style and para.style.name.startswith("Heading"):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif not paragraph_has_drawing(para):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for table in doc.tables:
        table.alignment = 1
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                if row_idx == 0:
                    shd = tc_pr.find(qn("w:shd"))
                    if shd is None:
                        shd = OxmlElement("w:shd")
                        tc_pr.append(shd)
                    shd.set(qn("w:fill"), "EDEDED")
                for p in cell.paragraphs:
                    p.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER
                        if row_idx == 0
                        else WD_ALIGN_PARAGRAPH.LEFT
                    )
                    for run in p.runs:
                        set_run_font(run, size=8.3, bold=True if row_idx == 0 else None)

    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size=7)

    doc.save(path)
    return path


if __name__ == "__main__":
    print(format_docx())
